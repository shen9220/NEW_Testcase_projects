import os
import tempfile
from fastapi import UploadFile


MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


class DocumentParseError(Exception):
    pass


async def parse_upload(file: UploadFile) -> tuple[str, str, bytes]:
    """
    Parse an uploaded PRD file into markdown text.
    Returns (filename, markdown_content, raw_bytes).
    """
    ext = os.path.splitext(file.filename or "untitled")[1].lower()
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise DocumentParseError(f"文件过大，最大支持 {MAX_FILE_SIZE // 1024 // 1024}MB")

    markdown = ""
    if ext == ".md":
        markdown = content.decode("utf-8", errors="replace")
    elif ext == ".docx":
        markdown = _parse_docx(content)
    elif ext == ".pdf":
        markdown = _parse_pdf(content)
    else:
        raise DocumentParseError(f"不支持的格式: {ext}，仅支持 .md / .docx / .pdf")

    if not markdown.strip():
        raise DocumentParseError("文件内容为空，无法解析")

    return file.filename or "untitled", markdown, content


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if "Heading 1" in style or style == "Title":
                lines.append(f"# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            else:
                lines.append(text)
        # Also extract tables
        for table in doc.tables:
            lines.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(lines)
    except Exception as e:
        raise DocumentParseError(f"Word 解析失败: {e}")


def _parse_pdf(content: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        lines = []
        for page in doc:
            text = page.get_text("markdown")
            if text:
                lines.append(text)
        doc.close()
        return "\n\n".join(lines)
    except Exception as e:
        raise DocumentParseError(f"PDF 解析失败: {e}")


def parse_markdown_text(content: str) -> str:
    """Validate and clean up manually entered markdown."""
    if not content.strip():
        raise DocumentParseError("内容为空")
    return content
